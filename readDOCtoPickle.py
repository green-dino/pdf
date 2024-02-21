import os
import pandas as pd
import docx
from docx import Document
import pickle
import re

class FileReader:
    @staticmethod
    def read_docx(file_path):
        try:
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs]
            return paragraphs
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return None

    @staticmethod
    def read_doc(file_path):
        try:
            with open(file_path, 'rb') as doc_file:
                doc = Document(doc_file)
                paragraphs = [para.text for para in doc.paragraphs]
                return paragraphs
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return None

class FileHandler:
    @staticmethod
    def list_files(directory):
        return [f for f in os.listdir(directory) if f.endswith(('.doc', '.docx'))]

    @staticmethod
    def create_output_directory(directory):
        output_dir = os.path.join(directory, "output")
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        return output_dir

class DataProcessor:
    def __init__(self):
        self.email_pattern = re.compile(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,4}')
        self.word_pattern = re.compile(r'\b[a-zA-Z]{5,15}\b')
        self.url_pattern = re.compile(r'\b\w+:\/\/[\w@][\w.:@]+\/?[\w.\.?=%&=\-@$,]*\b')

    def process_data(self, df):
        new_columns = {}
        for column in df.columns:
            if df[column].dtype == 'object':
                new_columns[column + '_emails'] = df[column].apply(lambda x: self.email_pattern.findall(str(x)))
                new_columns[column + '_words'] = df[column].apply(lambda x: self.word_pattern.findall(str(x)))
                new_columns[column + '_urls'] = df[column].apply(lambda x: self.url_pattern.findall(str(x)))
        return pd.concat([df, pd.DataFrame(new_columns)], axis=1)

class DataFrameHandler:
    @staticmethod
    @staticmethod
    def read_file_to_dataframe(file_path):
        try:
            if file_path.endswith('.csv'):
                return pd.read_csv(file_path)
            elif file_path.endswith(('.xls', '.xlsx')):
                return pd.read_excel(file_path)
            elif file_path.endswith('.docx'):
                doc = docx.Document(file_path)
                paragraphs = [para.text for para in doc.paragraphs]
                df = pd.DataFrame(paragraphs, columns=['Content'])
                return df.transpose()  # Transpose to create separate columns for each paragraph
            elif file_path.endswith('.doc'):
                with open(file_path, 'rb') as doc_file:
                    doc = Document(doc_file)
                    paragraphs = [para.text for para in doc.paragraphs]
                    df = pd.DataFrame(paragraphs, columns=['Content'])
                    return df.transpose()  # Transpose to create separate columns for each paragraph
            else:
                raise ValueError("Unsupported file format. Please provide a CSV, XLS/XLSX, DOCX, or DOC file.")
        except Exception as e:
            print(f"Error reading {file_path}: {e}")
            return None


    @staticmethod
    def save_dataframe(df, filename, format='csv'):
        try:
            if format == 'csv':
                df.to_csv(filename, index=False)
            elif format == 'pickle':
                df.to_pickle(filename)
            else:
                raise ValueError("Unsupported format. Please choose 'csv' or 'pickle'.")
            print(f"DataFrame saved as {format.upper()}: {filename}")
        except Exception as e:
            print(f"Error saving {format.upper()} file: {e}")

def get_directory():
    while True:
        file_dir = input("Enter the directory where the file is stored (type 'exit' to quit): ").strip()
        if file_dir.lower() == 'exit':
            break

        if not os.path.isdir(file_dir):
            print("Invalid directory path.")
            continue

        files = FileHandler.list_files(file_dir)
        if not files:
            print("No DOC or DOCX files found in the specified directory.")
            continue

        print("Files found in the directory:")
        for i, file in enumerate(files, start=1):
            print(f"{i}. {file}")

        yield file_dir, files

def process_file(file_dir, files):
    try:
        file_choice = int(input("Enter the number corresponding to the file you want to use: "))
        file_path = os.path.join(file_dir, files[file_choice - 1])
        return DataFrameHandler.read_file_to_dataframe(file_path)
    except Exception as e:
        print(f"Error processing file: {e}")
        return None

def save_dataframe(output_dir, dataframe):
    try:
        output_filename = input("Enter the filename for the output file (without extension): ")
        format_choice = input("Choose output format ('csv' or 'pickle'): ").lower()

        output_file = os.path.join(output_dir, f"{output_filename}.{format_choice}")
        DataFrameHandler.save_dataframe(dataframe, output_file, format_choice)
    except Exception as e:
        print(f"Error saving DataFrame: {e}")

if __name__ == "__main__":
    processor = DataProcessor()

    dir_gen = get_directory()
    for file_dir, files in dir_gen:
        output_dir = FileHandler.create_output_directory(file_dir)
        dataframe = process_file(file_dir, files)
        if dataframe is not None:
            save_dataframe(output_dir, dataframe)
